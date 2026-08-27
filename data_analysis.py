# data_analysis.py

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import os
import glob
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def analyze_data():
    """Run basic data analysis from Super Admin export"""
    print("=" * 70)
    print("📊 DATA ANALYSIS REPORT")
    print("=" * 70)
    
    # ===== Find the export file =====
    export_file = None
    search_locations = ['.', 'data_exports']
    
    print("🔍 Searching for export file...")
    
    for location in search_locations:
        if os.path.exists(location):
            files = glob.glob(os.path.join(location, 'corep_data_*.xlsx'))
            if files:
                export_file = max(files, key=os.path.getctime)
                print(f"  ✅ Found: {os.path.basename(export_file)} in '{location}'")
                break
    
    if not export_file:
        print("❌ No 'corep_data_*.xlsx' file found!")
        return
    
    # ===== Load the file =====
    print(f"\n📂 Loading: {os.path.basename(export_file)}")
    print(f"📅 File created: {datetime.fromtimestamp(os.path.getctime(export_file)).strftime('%Y-%m-%d %H:%M:%S')}")
    
    try:
        excel_file = pd.ExcelFile(export_file)
        sheet_names = excel_file.sheet_names
        print(f"📊 Found {len(sheet_names)} sheets: {sheet_names}")
        
        dataframes = {}
        for sheet in sheet_names:
            dataframes[sheet] = pd.read_excel(export_file, sheet_name=sheet)
            print(f"  ✅ Loaded sheet: {sheet} ({len(dataframes[sheet])} rows)")
        
        patients = dataframes.get('Patients', pd.DataFrame())
        nursing = dataframes.get('Nursing_Assessments', pd.DataFrame())
        consultations = dataframes.get('Consultations', pd.DataFrame())
        lab_tests = dataframes.get('Lab_Tests', pd.DataFrame())
        optical = dataframes.get('Optical_Assessments', pd.DataFrame())
        pharmacy = dataframes.get('Pharmacy_Orders', pd.DataFrame())
        drugs = dataframes.get('Drugs', pd.DataFrame())
        notifications = dataframes.get('Notifications', pd.DataFrame())
        dispensing = dataframes.get('Dispensing', pd.DataFrame())
        user_profiles = dataframes.get('User_Profiles', pd.DataFrame())
        
        print("✅ File loaded successfully!")
        
    except Exception as e:
        print(f"❌ Error loading file: {e}")
        return
    
    if patients.empty:
        print("❌ No patient data found!")
        return
    
    # ===== Normalize column names =====
    for df_name, df in [('patients', patients), ('consultations', consultations), 
                         ('lab_tests', lab_tests), ('optical', optical), 
                         ('pharmacy', pharmacy), ('drugs', drugs),
                         ('notifications', notifications), ('dispensing', dispensing),
                         ('user_profiles', user_profiles)]:
        if not df.empty:
            df.columns = [c.lower().replace(' ', '_') for c in df.columns]
    
    # ===== Run Analysis and Store Results =====
    results = {}
    
    # ===== PATIENT DEMOGRAPHICS =====
    print("\n" + "=" * 70)
    print("👥 PATIENT DEMOGRAPHICS")
    print("-" * 40)
    
    results['patients'] = {
        'total': len(patients),
        'gender': {},
        'age': {},
        'stages': {}
    }
    
    if 'gender' in patients.columns:
        gender_counts = patients['gender'].value_counts()
        results['patients']['gender'] = gender_counts.to_dict()
        print(f"Gender: {gender_counts.to_dict()}")
    
    if 'date_of_birth' in patients.columns:
        patients['date_of_birth'] = pd.to_datetime(patients['date_of_birth'])
        patients['age'] = (datetime.now() - patients['date_of_birth']).dt.days // 365
        results['patients']['age'] = {
            'average': round(patients['age'].mean(), 1),
            'min': patients['age'].min(),
            'max': patients['age'].max(),
            'median': patients['age'].median()
        }
        print(f"Average Age: {patients['age'].mean():.1f} years")
        print(f"Age Range: {patients['age'].min()} - {patients['age'].max()} years")
    
    if 'current_stage' in patients.columns:
        stage_counts = patients['current_stage'].value_counts()
        results['patients']['stages'] = stage_counts.to_dict()
        print(f"Stages: {stage_counts.to_dict()}")
    
    # ===== CONSULTATION ANALYSIS =====
    results['consultations'] = {}
    if not consultations.empty:
        print("\n" + "=" * 70)
        print("👨‍⚕️ CONSULTATION ANALYSIS")
        print("-" * 40)
        
        results['consultations']['total'] = len(consultations)
        print(f"Total Consultations: {len(consultations)}")
        
        if 'completed' in consultations.columns:
            results['consultations']['completed'] = int(consultations['completed'].sum())
            print(f"Completed: {consultations['completed'].sum()}")
        
        results['consultations']['referrals'] = {}
        print("Referrals:")
        for ref in ['pharmacy', 'laboratory', 'optician', 'specialist']:
            col = f'refer_to_{ref}'
            if col in consultations.columns:
                count = int(consultations[col].sum())
                results['consultations']['referrals'][ref] = count
                print(f"  {ref.capitalize()}: {count}")
        
        if 'diagnosis' in consultations.columns:
            top_diagnoses = consultations['diagnosis'].value_counts().head(5)
            results['consultations']['top_diagnoses'] = top_diagnoses.to_dict()
            print(f"Top Diagnoses: {top_diagnoses.to_dict()}")
    
    # ===== LAB TEST ANALYSIS =====
    results['lab_tests'] = {}
    if not lab_tests.empty:
        print("\n" + "=" * 70)
        print("🧪 LAB TEST ANALYSIS")
        print("-" * 40)
        
        results['lab_tests']['total'] = len(lab_tests)
        print(f"Total Tests: {len(lab_tests)}")
        
        if 'completed' in lab_tests.columns:
            results['lab_tests']['completed'] = int(lab_tests['completed'].sum())
            print(f"Completed: {lab_tests['completed'].sum()}")
        
        if 'malaria_parasite' in lab_tests.columns:
            malaria_counts = lab_tests['malaria_parasite'].value_counts()
            results['lab_tests']['malaria'] = malaria_counts.to_dict()
            print(f"Malaria: {malaria_counts.to_dict()}")
        
        if 'hbsag' in lab_tests.columns:
            hbsag_counts = lab_tests['hbsag'].value_counts()
            results['lab_tests']['hbsag'] = hbsag_counts.to_dict()
            print(f"HBsAg: {hbsag_counts.to_dict()}")
        
        if 'random_blood_sugar' in lab_tests.columns:
            rbs = lab_tests['random_blood_sugar'].dropna()
            if not rbs.empty:
                results['lab_tests']['rbs'] = {
                    'average': round(rbs.mean(), 1),
                    'min': rbs.min(),
                    'max': rbs.max()
                }
                print(f"Average RBS: {rbs.mean():.1f} mmol/L")
                print(f"RBS Range: {rbs.min():.1f} - {rbs.max():.1f} mmol/L")
    
    # ===== OPTICAL ANALYSIS =====
    results['optical'] = {}
    if not optical.empty:
        print("\n" + "=" * 70)
        print("👁️ OPTICAL ANALYSIS")
        print("-" * 40)
        
        results['optical']['total'] = len(optical)
        print(f"Total Assessments: {len(optical)}")
        
        if 'is_walk_in' in optical.columns:
            results['optical']['walk_ins'] = int(optical['is_walk_in'].sum())
            print(f"Walk-ins: {optical['is_walk_in'].sum()}")
        
        if 'glasses_allocated' in optical.columns:
            results['optical']['glasses_total'] = int(optical['glasses_allocated'].sum())
            results['optical']['glasses_avg'] = round(optical['glasses_allocated'].mean(), 1)
            print(f"Glasses Allocated: {optical['glasses_allocated'].sum()}")
            print(f"Average Glasses: {optical['glasses_allocated'].mean():.1f}")
        
        if 'visual_acuity_left' in optical.columns:
            va_left = optical['visual_acuity_left'].dropna()
            if not va_left.empty:
                results['optical']['va_left'] = va_left.value_counts().head(3).to_dict()
                print(f"Visual Acuity (Left): {va_left.value_counts().head(3).to_dict()}")
    
    # ===== PHARMACY ANALYSIS =====
    results['pharmacy'] = {}
    if not pharmacy.empty:
        print("\n" + "=" * 70)
        print("💊 PHARMACY ANALYSIS")
        print("-" * 40)
        
        results['pharmacy']['total'] = len(pharmacy)
        print(f"Total Orders: {len(pharmacy)}")
        
        if 'quantity' in pharmacy.columns:
            results['pharmacy']['total_quantity'] = int(pharmacy['quantity'].sum())
            results['pharmacy']['avg_quantity'] = round(pharmacy['quantity'].mean(), 1)
            print(f"Total Quantity: {pharmacy['quantity'].sum()}")
            print(f"Average Order: {pharmacy['quantity'].mean():.1f}")
        
        if 'dispensed' in pharmacy.columns:
            results['pharmacy']['dispensed'] = int(pharmacy['dispensed'].sum())
            print(f"Dispensed: {pharmacy['dispensed'].sum()}")
        
        if 'drug_name' in pharmacy.columns:
            top_drugs = pharmacy['drug_name'].value_counts().head(5)
            results['pharmacy']['top_drugs'] = top_drugs.to_dict()
            print(f"Top Drugs: {top_drugs.to_dict()}")
    
    # ===== DRUG INVENTORY =====
    results['drugs'] = {}
    if not drugs.empty:
        print("\n" + "=" * 70)
        print("💊 DRUG INVENTORY")
        print("-" * 40)
        
        results['drugs']['total'] = len(drugs)
        print(f"Total Drugs: {len(drugs)}")
        
        if 'quantity' in drugs.columns:
            results['drugs']['total_stock'] = int(drugs['quantity'].sum())
            low_stock = drugs[drugs['quantity'] < 50]
            results['drugs']['low_stock'] = len(low_stock)
            print(f"Total Stock: {drugs['quantity'].sum()}")
            print(f"Low Stock (<50): {len(low_stock)}")
        
        if 'category' in drugs.columns:
            results['drugs']['categories'] = drugs['category'].value_counts().to_dict()
            print(f"Categories: {drugs['category'].value_counts().to_dict()}")
    
    # ===== NOTIFICATIONS =====
    results['notifications'] = {}
    if not notifications.empty:
        print("\n" + "=" * 70)
        print("🔔 NOTIFICATIONS")
        print("-" * 40)
        
        results['notifications']['total'] = len(notifications)
        print(f"Total Notifications: {len(notifications)}")
        
        if 'is_read' in notifications.columns:
            results['notifications']['read'] = int(notifications['is_read'].sum())
            results['notifications']['unread'] = len(notifications) - int(notifications['is_read'].sum())
            print(f"Read: {notifications['is_read'].sum()}")
            print(f"Unread: {len(notifications) - notifications['is_read'].sum()}")
    
    # ===== DISPENSING =====
    results['dispensing'] = {}
    if not dispensing.empty:
        print("\n" + "=" * 70)
        print("💊 DISPENSING RECORDS")
        print("-" * 40)
        
        results['dispensing']['total'] = len(dispensing)
        print(f"Total Dispensing Records: {len(dispensing)}")
        
        if 'quantity_dispensed' in dispensing.columns:
            results['dispensing']['total_quantity'] = int(dispensing['quantity_dispensed'].sum())
            results['dispensing']['avg_quantity'] = round(dispensing['quantity_dispensed'].mean(), 1)
            print(f"Total Quantity Dispensed: {dispensing['quantity_dispensed'].sum()}")
            print(f"Average per Dispense: {dispensing['quantity_dispensed'].mean():.1f}")
    
    # ===== USER PROFILES =====
    results['users'] = {}
    if not user_profiles.empty:
        print("\n" + "=" * 70)
        print("👤 USER PROFILES")
        print("-" * 40)
        
        results['users']['total'] = len(user_profiles)
        print(f"Total Users: {len(user_profiles)}")
        
        if 'role' in user_profiles.columns:
            results['users']['roles'] = user_profiles['role'].value_counts().to_dict()
            print(f"Roles: {user_profiles['role'].value_counts().to_dict()}")
        
        if 'is_active' in user_profiles.columns:
            results['users']['active'] = int(user_profiles['is_active'].sum())
            results['users']['inactive'] = len(user_profiles) - int(user_profiles['is_active'].sum())
            print(f"Active Users: {user_profiles['is_active'].sum()}")
            print(f"Inactive Users: {len(user_profiles) - user_profiles['is_active'].sum()}")
    
    # ===== SUMMARY =====
    print("\n" + "=" * 70)
    print("📊 SUMMARY STATISTICS")
    print("-" * 40)
    print(f"Total Patients: {len(patients)}")
    print(f"Total Consultations: {len(consultations) if not consultations.empty else 0}")
    print(f"Total Lab Tests: {len(lab_tests) if not lab_tests.empty else 0}")
    print(f"Total Optical Assessments: {len(optical) if not optical.empty else 0}")
    print(f"Total Pharmacy Orders: {len(pharmacy) if not pharmacy.empty else 0}")
    print(f"Total Drugs: {len(drugs) if not drugs.empty else 0}")
    print(f"Total Notifications: {len(notifications) if not notifications.empty else 0}")
    print(f"Total Dispensing Records: {len(dispensing) if not dispensing.empty else 0}")
    print(f"Total Users: {len(user_profiles) if not user_profiles.empty else 0}")
    
    print("\n" + "=" * 70)
    print("✅ Analysis Complete!")
    
    # ===== Create Charts =====
    try:
        chart_files = create_charts(patients, consultations, lab_tests, optical, pharmacy, drugs)
    except Exception as e:
        print(f"⚠️ Charts could not be created: {e}")
        chart_files = {}
    
    # ===== Create Word Report =====
    try:
        create_word_report(results, chart_files, export_file)
    except Exception as e:
        print(f"⚠️ Word report could not be created: {e}")
        print("Make sure python-docx is installed: pip install python-docx")

def create_charts(patients, consultations, lab_tests, optical, pharmacy, drugs):
    """Create basic charts for analysis"""
    print("\n📈 Creating charts...")
    
    charts_dir = 'data_charts'
    if not os.path.exists(charts_dir):
        os.makedirs(charts_dir)
    
    chart_files = {}
    chart_count = 0
    
    # 1. Gender distribution
    if not patients.empty and 'gender' in patients.columns:
        plt.figure(figsize=(8, 6))
        patients['gender'].value_counts().plot(kind='pie', autopct='%1.1f%%')
        plt.title('Patient Gender Distribution')
        plt.ylabel('')
        plt.tight_layout()
        filename = f'{charts_dir}/gender_distribution.png'
        plt.savefig(filename)
        plt.close()
        chart_files['gender'] = filename
        chart_count += 1
        print(f"  ✅ Gender chart saved")
    
    # 2. Age distribution
    if not patients.empty and 'age' in patients.columns:
        plt.figure(figsize=(10, 6))
        patients['age'].hist(bins=20, edgecolor='black')
        plt.title('Patient Age Distribution')
        plt.xlabel('Age (years)')
        plt.ylabel('Number of Patients')
        plt.tight_layout()
        filename = f'{charts_dir}/age_distribution.png'
        plt.savefig(filename)
        plt.close()
        chart_files['age'] = filename
        chart_count += 1
        print(f"  ✅ Age chart saved")
    
    # 3. Consultation referrals
    if not consultations.empty:
        referral_cols = ['refer_to_pharmacy', 'refer_to_laboratory', 'refer_to_optician', 'refer_to_specialist']
        referral_data = {}
        for col in referral_cols:
            if col in consultations.columns:
                referral_data[col.replace('refer_to_', '').capitalize()] = consultations[col].sum()
        
        if referral_data:
            plt.figure(figsize=(8, 6))
            pd.Series(referral_data).plot(kind='bar', color=['#2e7d32', '#1a73e8', '#6a1b9a', '#f57c00'])
            plt.title('Consultation Referrals')
            plt.xlabel('Referral Type')
            plt.ylabel('Count')
            plt.tight_layout()
            filename = f'{charts_dir}/referrals.png'
            plt.savefig(filename)
            plt.close()
            chart_files['referrals'] = filename
            chart_count += 1
            print(f"  ✅ Referrals chart saved")
    
    # 4. Lab test results
    if not lab_tests.empty and 'malaria_parasite' in lab_tests.columns:
        plt.figure(figsize=(8, 6))
        lab_tests['malaria_parasite'].value_counts().plot(kind='bar', color=['#c62828', '#2e7d32', '#f57c00'])
        plt.title('Malaria Parasite Results')
        plt.xlabel('Result')
        plt.ylabel('Count')
        plt.tight_layout()
        filename = f'{charts_dir}/malaria_results.png'
        plt.savefig(filename)
        plt.close()
        chart_files['malaria'] = filename
        chart_count += 1
        print(f"  ✅ Malaria results chart saved")
    
    # 5. Drug categories
    if not drugs.empty and 'category' in drugs.columns:
        plt.figure(figsize=(10, 6))
        drugs['category'].value_counts().plot(kind='bar')
        plt.title('Drug Categories')
        plt.xlabel('Category')
        plt.ylabel('Count')
        plt.tight_layout()
        filename = f'{charts_dir}/drug_categories.png'
        plt.savefig(filename)
        plt.close()
        chart_files['drug_categories'] = filename
        chart_count += 1
        print(f"  ✅ Drug categories chart saved")
    
    print(f"✅ {chart_count} charts saved to '{charts_dir}/' directory")
    return chart_files

def create_word_report(results, chart_files, export_file):
    """Create a Word document report"""
    print("\n📄 Creating Word Report...")
    
    doc = Document()
    
    # ===== Title =====
    title = doc.add_heading('COREP OUTREACH - DATA ANALYSIS REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== Date =====
    date_para = doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== Source =====
    doc.add_paragraph(f'Source File: {os.path.basename(export_file)}')
    doc.add_paragraph('=' * 70)
    
    # ===== Summary Section =====
    doc.add_heading('EXECUTIVE SUMMARY', level=1)
    
    # Summary table
    summary_data = [
        ['Metric', 'Count'],
        ['Total Patients', str(results.get('patients', {}).get('total', 0))],
        ['Total Consultations', str(results.get('consultations', {}).get('total', 0))],
        ['Total Lab Tests', str(results.get('lab_tests', {}).get('total', 0))],
        ['Total Optical Assessments', str(results.get('optical', {}).get('total', 0))],
        ['Total Pharmacy Orders', str(results.get('pharmacy', {}).get('total', 0))],
        ['Total Drugs', str(results.get('drugs', {}).get('total', 0))],
        ['Total Notifications', str(results.get('notifications', {}).get('total', 0))],
        ['Total Dispensing Records', str(results.get('dispensing', {}).get('total', 0))],
        ['Total Users', str(results.get('users', {}).get('total', 0))],
    ]
    
    table = doc.add_table(rows=len(summary_data), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(summary_data):
        row = table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        # Bold the first row
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # ===== Patient Demographics =====
    doc.add_heading('1. PATIENT DEMOGRAPHICS', level=1)
    
    patients_data = results.get('patients', {})
    doc.add_paragraph(f'Total Patients: {patients_data.get("total", 0)}')
    
    # Gender
    if patients_data.get('gender'):
        doc.add_heading('Gender Distribution:', level=2)
        for gender, count in patients_data['gender'].items():
            doc.add_paragraph(f'  • {gender}: {count}', style='List Bullet')
    
    # Age
    if patients_data.get('age'):
        doc.add_heading('Age Statistics:', level=2)
        age = patients_data['age']
        doc.add_paragraph(f'  • Average Age: {age.get("average", 0)} years')
        doc.add_paragraph(f'  • Median Age: {age.get("median", 0)} years')
        doc.add_paragraph(f'  • Age Range: {age.get("min", 0)} - {age.get("max", 0)} years')
    
    # Stages
    if patients_data.get('stages'):
        doc.add_heading('Workflow Stages:', level=2)
        for stage, count in patients_data['stages'].items():
            doc.add_paragraph(f'  • {stage}: {count}', style='List Bullet')
    
    doc.add_paragraph()
    
    # ===== Consultation Analysis =====
    doc.add_heading('2. CONSULTATION ANALYSIS', level=1)
    
    cons_data = results.get('consultations', {})
    doc.add_paragraph(f'Total Consultations: {cons_data.get("total", 0)}')
    if cons_data.get('completed') is not None:
        doc.add_paragraph(f'Completed Consultations: {cons_data.get("completed", 0)}')
    
    # Referrals
    if cons_data.get('referrals'):
        doc.add_heading('Referrals:', level=2)
        for ref, count in cons_data['referrals'].items():
            doc.add_paragraph(f'  • {ref.capitalize()}: {count}', style='List Bullet')
    
    # Top Diagnoses
    if cons_data.get('top_diagnoses'):
        doc.add_heading('Top 5 Diagnoses:', level=2)
        for diagnosis, count in cons_data['top_diagnoses'].items():
            doc.add_paragraph(f'  • {diagnosis}: {count}', style='List Bullet')
    
    doc.add_paragraph()
    
    # ===== Lab Test Analysis =====
    doc.add_heading('3. LABORATORY TEST ANALYSIS', level=1)
    
    lab_data = results.get('lab_tests', {})
    doc.add_paragraph(f'Total Lab Tests: {lab_data.get("total", 0)}')
    if lab_data.get('completed') is not None:
        doc.add_paragraph(f'Completed Tests: {lab_data.get("completed", 0)}')
    
    if lab_data.get('malaria'):
        doc.add_heading('Malaria Parasite Results:', level=2)
        for result, count in lab_data['malaria'].items():
            doc.add_paragraph(f'  • {result}: {count}', style='List Bullet')
    
    if lab_data.get('hbsag'):
        doc.add_heading('HBsAg Results:', level=2)
        for result, count in lab_data['hbsag'].items():
            doc.add_paragraph(f'  • {result}: {count}', style='List Bullet')
    
    if lab_data.get('rbs'):
        doc.add_heading('Random Blood Sugar:', level=2)
        rbs = lab_data['rbs']
        doc.add_paragraph(f'  • Average: {rbs.get("average", 0)} mmol/L')
        doc.add_paragraph(f'  • Range: {rbs.get("min", 0)} - {rbs.get("max", 0)} mmol/L')
    
    doc.add_paragraph()
    
    # ===== Optical Analysis =====
    doc.add_heading('4. OPTICAL ASSESSMENT ANALYSIS', level=1)
    
    optical_data = results.get('optical', {})
    doc.add_paragraph(f'Total Assessments: {optical_data.get("total", 0)}')
    if optical_data.get('walk_ins') is not None:
        doc.add_paragraph(f'Walk-in Patients: {optical_data.get("walk_ins", 0)}')
    if optical_data.get('glasses_total') is not None:
        doc.add_paragraph(f'Total Glasses Allocated: {optical_data.get("glasses_total", 0)}')
        doc.add_paragraph(f'Average Glasses per Patient: {optical_data.get("glasses_avg", 0)}')
    
    if optical_data.get('va_left'):
        doc.add_heading('Visual Acuity (Left Eye):', level=2)
        for va, count in optical_data['va_left'].items():
            doc.add_paragraph(f'  • {va}: {count}', style='List Bullet')
    
    doc.add_paragraph()
    
    # ===== Pharmacy Analysis =====
    doc.add_heading('5. PHARMACY ANALYSIS', level=1)
    
    pharm_data = results.get('pharmacy', {})
    doc.add_paragraph(f'Total Orders: {pharm_data.get("total", 0)}')
    if pharm_data.get('total_quantity') is not None:
        doc.add_paragraph(f'Total Quantity Dispensed: {pharm_data.get("total_quantity", 0)}')
        doc.add_paragraph(f'Average Order Size: {pharm_data.get("avg_quantity", 0)}')
    if pharm_data.get('dispensed') is not None:
        doc.add_paragraph(f'Orders Dispensed: {pharm_data.get("dispensed", 0)}')
    
    if pharm_data.get('top_drugs'):
        doc.add_heading('Top 5 Prescribed Drugs:', level=2)
        for drug, count in pharm_data['top_drugs'].items():
            doc.add_paragraph(f'  • {drug}: {count}', style='List Bullet')
    
    doc.add_paragraph()
    
    # ===== Drug Inventory =====
    doc.add_heading('6. DRUG INVENTORY', level=1)
    
    drugs_data = results.get('drugs', {})
    doc.add_paragraph(f'Total Drugs: {drugs_data.get("total", 0)}')
    if drugs_data.get('total_stock') is not None:
        doc.add_paragraph(f'Total Stock Units: {drugs_data.get("total_stock", 0)}')
        doc.add_paragraph(f'Low Stock Items (<50): {drugs_data.get("low_stock", 0)}')
    
    if drugs_data.get('categories'):
        doc.add_heading('Drug Categories:', level=2)
        for category, count in drugs_data['categories'].items():
            doc.add_paragraph(f'  • {category}: {count}', style='List Bullet')
    
    doc.add_paragraph()
    
    # ===== User Profiles =====
    if results.get('users'):
        doc.add_heading('7. USER PROFILES', level=1)
        user_data = results['users']
        doc.add_paragraph(f'Total Users: {user_data.get("total", 0)}')
        if user_data.get('active') is not None:
            doc.add_paragraph(f'Active Users: {user_data.get("active", 0)}')
            doc.add_paragraph(f'Inactive Users: {user_data.get("inactive", 0)}')
        
        if user_data.get('roles'):
            doc.add_heading('User Roles:', level=2)
            for role, count in user_data['roles'].items():
                doc.add_paragraph(f'  • {role}: {count}', style='List Bullet')
    
    doc.add_paragraph()
    
    # ===== Charts =====
    if chart_files:
        doc.add_heading('8. CHARTS', level=1)
        
        chart_order = ['gender', 'age', 'referrals', 'malaria', 'drug_categories']
        for chart_key in chart_order:
            if chart_key in chart_files:
                doc.add_heading(chart_key.replace('_', ' ').title(), level=2)
                try:
                    doc.add_picture(chart_files[chart_key], width=Inches(5.5))
                    doc.add_paragraph()
                except Exception as e:
                    doc.add_paragraph(f'(Chart image could not be embedded: {e})')
    
    # ===== Footer =====
    doc.add_paragraph('=' * 70)
    doc.add_paragraph(f'Report generated by COREP Outreach System')
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # ===== Save the document =====
    report_dir = 'data_reports'
    if not os.path.exists(report_dir):
        os.makedirs(report_dir)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'{report_dir}/corep_analysis_report_{timestamp}.docx'
    doc.save(filename)
    
    print(f"✅ Word Report saved: {filename}")
    return filename

if __name__ == "__main__":
    analyze_data()